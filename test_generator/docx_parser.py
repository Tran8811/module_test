# docx_parser.py
"""
Đọc file .docx và trả về:
  - lines: list[str] các dòng nội dung THEO ĐÚNG THỨ TỰ xuất hiện trong văn bản
           (bảng được thay bằng placeholder <tab>id</tab>, giống cơ chế production)
  - tables: dict[str, str]  id -> HTML string của bảng

Cần cài: pip install python-docx

Lưu ý so với bản production gốc:
  - Không trích ảnh (<image>id</image>) — nếu cần, có thể mở rộng thêm bằng
    cách duyệt document.inline_shapes, tương tự cách xử lý bảng bên dưới.
  - Dùng python-docx thuần, không qua LLM/OCR, nên chỉ đọc được các file
    .docx dạng text thật (không phải scan).
"""
import re
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


def _iter_block_items(doc: Document):
    """Duyệt các phần tử (đoạn văn, bảng) theo ĐÚNG thứ tự xuất hiện trong body.

    python-docx mặc định tách riêng doc.paragraphs và doc.tables (mất thứ tự
    xen kẽ giữa chúng) -> phải tự duyệt XML tree để giữ đúng thứ tự đọc.
    """
    parent_elm = doc.element.body
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _table_to_html(table: Table) -> str:
    rows_html = []
    for row in table.rows:
        cells_html = "".join(
            f"<td>{re.sub(r'<[^>]+>', '', cell.text).strip()}</td>" for cell in row.cells
        )
        rows_html.append(f"<tr>{cells_html}</tr>")
    return "<table>" + "".join(rows_html) + "</table>"


def extract_docx_lines(path: str) -> tuple[list[str], dict[str, str]]:
    doc = Document(path)

    lines: list[str] = []
    tables: dict[str, str] = {}
    table_counter = 0

    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        elif isinstance(block, Table):
            tid = f"t{table_counter}"
            tables[tid] = _table_to_html(block)
            lines.append(f"<tab>{tid}</tab>")
            table_counter += 1

    return lines, tables
